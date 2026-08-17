"""Generate the MS-thesis literature review chapter as an APA-7 formatted .docx.

Revision 3: retains the author's own Sections 2.1 and 2.2 verbatim and continues
the chapter in the same voice. All cited works are from 2020 onward (with the
author's original citations preserved). Every reference was verified against
publisher/indexing records (DOI, journal, volume, issue, pages) before inclusion.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"

doc = Document()

# ---------------------------------------------------------------- base style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)
pf.space_before = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ------------------------------------------------------- page number header
header_p = doc.sections[0].header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header_p.add_run()
fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
run.font.name = FONT
run.font.size = Pt(12)


# ------------------------------------------------------------------ helpers
def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    return p


def ref(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    return p


def page_break():
    doc.add_page_break()


# ================================================================= CHAPTER 2
h1("2. Literature review")

# --------------------------------------------------- 2.1 (author's own text)
h2("2.1 Concept of rainfall-runoff modeling")
body(
    "Rainfall\u2013runoff modeling is an important component of hydrological science because it provides a quantitative "
    "framework for understanding how precipitation is transformed into runoff within a watershed. The relationship between "
    "rainfall and runoff is controlled by several factors, including rainfall intensity and duration, soil characteristics, "
    "antecedent moisture conditions, land use and land cover, topography, drainage characteristics, evapotranspiration, and "
    "surface-water storage. Because these factors vary spatially and temporally, the response of a watershed to a given "
    "rainfall event is rarely linear or uniform. Hydrological models are therefore widely used to simplify and represent "
    "these processes and to estimate streamflow from available meteorological and watershed information (Sahu et al., 2023)."
)
body(
    "Rainfall\u2013runoff models have become increasingly important for flood forecasting, watershed management, "
    "water-resources planning and assessment of hydrological hazards. Recent research has produced a wide range of "
    "empirical and conceptual based data-driven models. However, model selection depends strongly on the purpose of the "
    "study, the characteristics of the watershed, the temporal and spatial resolution of available data, and the level of "
    "process representation required. Greater model complexity does not necessarily result in better performance because "
    "model uncertainty can arise from input data, parameterization, structural assumptions, and spatial heterogeneity "
    "(Sahu et al., 2023)."
)
body(
    "Among the available rainfall\u2013runoff models, the Hydrologic Engineering Center\u2013Hydrologic Modeling System "
    "(HEC-HMS) has been widely applied because of its flexible and modular structure. The model allows precipitation to be "
    "converted into runoff through different loss, transform, baseflow, and routing methods. An appropriate representation "
    "can be selected according to the characteristics and objectives of a particular watershed. Recent reviews have "
    "identified the Soil Conservation Service Curve Number (SCS-CN) and SCS Unit Hydrograph methods as among the most "
    "frequently used approaches in HEC-HMS applications, while recession methods and Muskingum-type routing approaches are "
    "also commonly applied (Labade et al., 2025; Sahu et al., 2023; Turkar et al., 2025)."
)
body(
    "The integration of HEC-HMS with Geographic Information Systems (GIS) has further increased its applicability. GIS "
    "enables the extraction of watershed boundaries, drainage networks, slopes, flow paths, land-use information, and soil "
    "characteristics from spatial datasets. These data can subsequently be used to develop hydrological parameters for the "
    "HEC-HMS model. Integrated ArcGIS-derived land-use, soil, and slope information with rainfall and runoff observations "
    "are used to develop HEC-HMS models for the Punpun River Basin (Ranjan & Singh, 2022). Their study demonstrated the "
    "practical value of combining spatial watershed information with observed hydrometeorological data for "
    "rainfall\u2013runoff simulation."
)
body(
    "The application of HEC-HMS is particularly relevant to Bangladesh because of the country\u2019s monsoonal rainfall "
    "regime, extensive river network, low-lying floodplains, wetlands, and recurrent flooding. Several recent studies have "
    "applied HEC-HMS to Bangladesh river basins, including the Khowai, Bangali, Halda, Old Brahmaputra, Dhaka, and "
    "Atrai\u2013Karatoa systems. These studies demonstrate the applicability of the model under Bangladesh\u2019s "
    "hydrological conditions but also highlight challenges related to rainfall representation, discharge availability, "
    "baseflow simulation, and model parameterization (M. B. Haque et al., 2024; Moniruzzaman & Mahalder, 2026; Nur et al., "
    "2022; Zhang et al., 2022)."
)
body(
    "The present study focuses on rainfall\u2013runoff modelling of the Gumai Beel\u2013Ichamoti River system using "
    "HEC-HMS. Gumai Beel represents a low-lying wetland and agricultural environment, while the Ichamoti River forms an "
    "important drainage component of the surrounding landscape. The hydrological response of such a system may differ from "
    "the conventional upland watersheds. Because wetland storage, seasonal inundation, drainage connectivity, and "
    "agricultural land use can influence the timing and magnitude of runoff. Therefore, a review of rainfall\u2013runoff "
    "processes, hydrological modelling approaches, HEC-HMS methodology, GIS integration, wetland hydrology, and previous "
    "Bangladesh applications is necessary to establish the scientific basis for the present research."
)

# --------------------------------------------------- 2.2 (author's own text)
h2("2.2 Rainfall-runoff process")
body(
    "Rainfall\u2013runoff transformation governs the hydrological routing of precipitation across a drainage basin and its "
    "ultimate conversion into streamflow discharge. Incident precipitation across the land surface is partitioned "
    "dynamically among canopy interception, infiltration into the vadose zone, surface depression storage, "
    "evapotranspiration loss, and overland or subsurface lateral flow. The volumetric fraction of precipitation that "
    "reaches the channel network is dictated by the dynamic equilibrium among these constituent processes. Consequently, "
    "the hydrograph configuration produced by a given precipitation event is governed not only by gross rainfall depth, but "
    "also by storm duration, temporal hyetograph distribution, peak rainfall intensity, and the antecedent hydrological "
    "state of the catchment (Labade et al., 2025)."
)
body(
    "Precipitation intensity directly regulates the threshold between infiltration and direct overland flow. When "
    "precipitation rates surpass the saturated hydraulic conductivity and infiltration capacity of the soil, "
    "infiltration-excess (Hortonian) overland flow is initiated and rapidly routed to the channel network. Conversely, "
    "rainfall rates below the infiltration threshold primarily replenish soil moisture storage or contribute to subsurface "
    "matrix flow. Antecedent moisture conditions further modulate this response: a dry catchment exhibits substantial "
    "initial abstraction and retention capacity, whereas antecedent saturation markedly reduces soil matrix suction "
    "potential, accelerating saturation-excess overland flow. As a result, identical storm depths falling under differing "
    "initial moisture conditions generate fundamentally disparate runoff hydrographs (Shi & Wang, 2020)."
)
body(
    "Catchment physiography and surface cover impose critical boundary controls on hydrograph translation and attenuation. "
    "Parameters such as soil texture, effective porosity, hydraulic conductivity, soil profile depth, topographic gradient, "
    "drainage density, and surface roughness collectively determine the velocity and storage dynamics of routed water. "
    "Anthropogenic land-use and land-cover (LULC) modifications directly influence these hydrologic controls such as urban "
    "expansion increases surface imperviousness and accelerates hydrograph. Dense vegetative canopies enhance interception "
    "and infiltration capacity. In agricultural floodplains, hydrological responses exhibit pronounced seasonal variability "
    "driven by cyclical crop growth stages, tillage practices, and artificial field drainage regimes (Dibaba et al., 2020)."
)
body(
    "In low-lying wetland such as Gumai Beel, runoff generation is influenced by dynamic depression storage, delayed "
    "infiltration, and temporary surface-water retention. Within the floodplain morphology of Bangladesh, river\u2013wetland "
    "systems function as hydrodynamic buffers whose storage behavior is dictated by seasonal monsoon inundation pulses and "
    "dynamic lateral river connectivity (Adnan et al., 2020; Akter & Sawon, 2024). These seasonal fluctuations in backwater "
    "staging, storage volume, and hydraulic connectivity introduce significant attenuation and lag into the catchment "
    "discharge hydrograph (Akter & Sawon, 2024). It represents a critical hydrodynamic mechanism in the rainfall\u2013runoff "
    "modeling of the Gumai Beel\u2013Ichamoti River system."
)

# =============================================== 2.3 (continuation begins) ==
h2("2.3 Classification and selection of hydrological models")
body(
    "Hydrological models are conventionally differentiated along two principal axes: the degree of physical process "
    "representation and the level of spatial discretisation. On the first axis, empirical or data-driven models establish "
    "statistical transfer functions between input and output series without explicit reference to physical mechanisms; "
    "conceptual models represent the catchment as a cascade of interconnected storage elements whose governing equations "
    "approximate physical behaviour; and physically based models solve conservation equations for mass, momentum, and "
    "energy using parameters that are, in principle, independently measurable. On the second axis, lumped models treat the "
    "catchment as a single homogeneous unit, semi-distributed models subdivide it into sub-basins within which parameters "
    "are averaged, and fully distributed models resolve hydrological state variables on a regular computational grid "
    "(Sahu et al., 2023). HEC-HMS occupies an intermediate position in this taxonomy, functioning as a conceptual, "
    "semi-distributed system that reproduces the dominant runoff-generation mechanisms without imposing the data burden of "
    "fully distributed formulations (Turkar et al., 2025)."
)
body(
    "The proliferation of hydrological models has itself become a subject of critical enquiry. Horton et al. (2022), "
    "combining a systematic literature review with a structured survey of practising modellers, demonstrated that model "
    "diversity is driven only partially by genuine differences in application context; a substantial component reflects "
    "institutional familiarity, code accessibility, and the reuse of pre-existing model configurations. Their central "
    "observation, that the justification for model selection is rarely articulated in published studies, establishes an "
    "explicit methodological obligation: the choice of a modelling system must be defended with reference to the physical "
    "characteristics of the catchment, the availability and resolution of input data, and the specific objectives of the "
    "investigation. This obligation is addressed directly in the present study through the reasoning developed in "
    "Sections 2.3 and 2.4."
)
body(
    "For data-limited catchments, the operative comparison in contemporary practice lies between conceptual "
    "semi-distributed systems such as HEC-HMS and more heavily parameterised ecohydrological frameworks such as the Soil "
    "and Water Assessment Tool (SWAT). Comparative evaluations clarify the associated trade-offs. Aliye et al. (2020), "
    "examining both models within a data-scarce catchment of the Ethiopian Rift Valley Lakes Basin, concluded that each "
    "framework was capable of reproducing observed streamflow with acceptable accuracy. Moniruzzaman and Mahalder (2026), "
    "applying an equivalent comparison to the Atrai\u2013Karatoa River Basin of northern Bangladesh, reported that SWAT "
    "reproduced high flows more faithfully whereas HEC-HMS achieved superior accuracy across the medium-flow range, with "
    "both models simulating low flows adequately. These findings indicate that HEC-HMS constitutes a defensible selection "
    "where the modelling objective is hydrograph generation rather than the simulation of sediment transport, nutrient "
    "flux, or land-management scenarios, and where the parsimony of its parameter set corresponds to the available "
    "observational record."
)
body(
    "A further operational distinction separates event-based from continuous simulation. Event-based configurations "
    "reproduce the catchment response to discrete storm episodes and are appropriate for design-flood estimation, peak "
    "discharge analysis, and flood forecasting. Continuous configurations track soil-moisture accounting through "
    "alternating wet and dry periods over multi-annual horizons and are appropriate for water-balance and water-resources "
    "assessment. Reviews of the HEC-HMS application literature indicate that the SCS-CN loss method dominates event-based "
    "practice, whereas the Soil Moisture Accounting (SMA) algorithm is preferred for continuous simulation (Labade et al., "
    "2025; Turkar et al., 2025). Because the inundation regime of the Gumai Beel is governed simultaneously by discrete "
    "monsoon storm events and by the strongly seasonal water balance of the surrounding floodplain, the present study "
    "requires a modelling system capable of operating in both modes, a requirement that HEC-HMS satisfies within a single "
    "computational framework."
)

h2("2.4 Structure and components of the HEC-HMS model")
body(
    "HEC-HMS was developed by the Hydrologic Engineering Center of the United States Army Corps of Engineers to simulate "
    "the complete precipitation\u2013runoff process of dendritic watershed systems (U.S. Army Corps of Engineers [USACE], "
    "n.d.). A model project comprises three principal components. The basin model represents the physical catchment as a "
    "network of hydrological elements, including sub-basins, reaches, junctions, reservoirs, diversions, sources, and "
    "sinks. The meteorological model assigns precipitation and evapotranspiration boundary conditions to each sub-basin "
    "through gauge weighting, inverse-distance interpolation, or gridded input. The control specifications define the "
    "temporal window and computational time step of the simulation. Within each sub-basin, the user selects one algorithm "
    "from each of several interchangeable method libraries governing canopy and surface storage, infiltration loss, "
    "direct-runoff transformation, and baseflow, while each routing reach is assigned an independent channel-routing "
    "method (Sahu et al., 2023; USACE, n.d.). This modular architecture accounts for the model\u2019s adaptability across "
    "climatic regimes and explains its widespread adoption in regions where proprietary modelling systems are financially "
    "inaccessible (Hamdan et al., 2021)."
)

h3("2.4.1 Loss methods")
body(
    "The loss model determines the proportion of incident precipitation abstracted through infiltration, interception, and "
    "surface retention, and therefore the residual depth of precipitation excess available for direct runoff. The SCS-CN "
    "method remains the dominant formulation within the HEC-HMS application literature because it consolidates the "
    "combined influence of hydrologic soil group, land use and treatment, surface condition, and antecedent moisture into "
    "a single dimensionless parameter that can be derived from mappable catchment attributes (Labade et al., 2025; "
    "Soulis, 2021). Soulis (2021), reviewing the status of the method after six decades of application, attributed its "
    "persistence to its conceptual simplicity, well-documented input requirements, and transferability to ungauged "
    "catchments, while identifying several unresolved limitations that remain active research topics. Foremost among these "
    "is the initial abstraction ratio: numerous investigations reviewed in that work reported locally calibrated values "
    "substantially below the conventional figure of 0.2, frequently approaching 0.05 or lower, implying systematic "
    "overestimation of initial losses when the standard ratio is applied uncritically."
)
body(
    "Complementary research has sought to extend the method by incorporating variables absent from its original "
    "formulation. Shi and Wang (2020) developed a modified SCS-CN equation combining the tabulated curve number with "
    "explicit slope-gradient, soil-moisture, and storm-duration factors, and demonstrated that the modified formulation "
    "raised model efficiency to approximately 80% during both calibration and validation while reducing root-mean-square "
    "error from 5.53 mm to 2.01 mm relative to the standard method. Their sensitivity analysis identified soil-moisture "
    "parameters as the most influential, followed by storm duration and slope, with the initial abstraction ratio the "
    "least sensitive. These results carry direct methodological implications for the present study, in which antecedent "
    "moisture in a seasonally saturated beel environment varies far more widely than in the agricultural watersheds from "
    "which the original curve-number tables were derived."
)
body(
    "HEC-HMS provides several alternatives to the curve-number approach, including the initial-and-constant and "
    "deficit-and-constant loss methods, the Green\u2013Ampt infiltration model, and the multi-layer SMA algorithm intended "
    "for continuous simulation (USACE, n.d.). Empirical evidence indicates that the SCS-CN method performs reliably in "
    "monsoonal and semi-arid event simulation, as demonstrated by Hamdan et al. (2021) for the Al-Adhaim catchment of "
    "northern Iraq and by Ranjan and Singh (2022) for the Punpun River Basin of eastern India. For multi-annual continuous "
    "applications, the SMA formulation is generally preferred, as adopted by Jawad (2024) in the Brahmaputra Basin. A "
    "consistent finding across this literature is that the curve number constitutes the single most sensitive parameter of "
    "the model. Turkar et al. (2025) identified curve number, infiltration rate, lag time, and baseflow as the controlling "
    "parameters across the studies they reviewed, and M. B. Haque et al. (2024) reported that simulated runoff in the "
    "Halda catchment was highly sensitive to sub-basin curve numbers, which depend directly on the accuracy of the "
    "underlying land-cover classification."
)

h3("2.4.2 Transform methods")
body(
    "The transform model converts the precipitation excess computed by the loss model into a direct-runoff hydrograph at "
    "the sub-basin outlet. The majority of applications employ unit-hydrograph theory, for which HEC-HMS provides the SCS "
    "dimensionless unit hydrograph, parameterised by basin lag time; the Clark unit hydrograph, which represents "
    "translation and attenuation separately through a time\u2013area histogram and a linear-reservoir storage coefficient; "
    "and the Snyder synthetic unit hydrograph, parameterised by lag and peaking coefficients (USACE, n.d.). Review "
    "evidence establishes the SCS unit hydrograph as the most frequently adopted transform across international practice, "
    "typically paired with SCS-CN losses (Labade et al., 2025; Sahu et al., 2023), a configuration applied successfully by "
    "Hamdan et al. (2021), Ranjan and Singh (2022), and Nujhat et al. (2024). The Clark formulation has been preferred in "
    "several large-basin continuous applications, including the Brahmaputra model of Jawad (2024), where storage "
    "coefficient and time of concentration were treated as primary calibration parameters."
)
body(
    "For extremely flat catchments, the estimation of lag time warrants particular caution. Empirical lag equations were "
    "predominantly derived from sloping terrain, and any misestimation propagates directly into errors in the timing of "
    "the simulated hydrograph peak. In floodplain environments where topographic gradients approach the vertical "
    "resolution of available elevation data, calibration of lag and storage parameters against observed hydrographs "
    "becomes indispensable rather than optional, an approach reflected in the optimisation strategies reported by Jawad "
    "(2024) and Admas et al. (2025)."
)

h3("2.4.3 Baseflow methods")
body(
    "Baseflow representation constitutes the most frequently underperforming component of event-oriented hydrological "
    "models, yet in floodplain-wetland systems the slow-drainage component may dominate both the recession limb of the "
    "hydrograph and the dry-season water balance. HEC-HMS offers recession, bounded-recession, linear-reservoir, and "
    "constant-monthly baseflow formulations (USACE, n.d.). Review evidence indicates that the recession method is the most "
    "widely adopted in international practice owing to its stability and its realistic representation of groundwater "
    "depletion (Labade et al., 2025), whereas the linear-reservoir formulation with multiple groundwater layers has been "
    "preferred in continuous large-basin applications (Jawad, 2024)."
)
body(
    "Recent Bangladeshi experience indicates that baseflow simulation represents a genuine methodological difficulty "
    "rather than a routine calibration exercise. M. B. Haque et al. (2024), modelling the Halda River catchment, obtained "
    "satisfactory aggregate performance statistics but reported a poor correspondence between observed and simulated flows "
    "along the baseflow portion of the hydrograph during calibration, attributing the discrepancy to unrepresented "
    "groundwater\u2013surface water exchange and recommending offline coupling with a groundwater model as a route to "
    "improvement. Raihan et al. (2020) similarly identified groundwater delay time and the baseflow alpha factor among the "
    "most sensitive parameters governing simulation quality in the Upper Halda Basin. These findings caution against "
    "interpreting event-calibrated models as complete descriptions of low-flow behaviour, a caution of particular force in "
    "beel environments where monsoon storage is released gradually throughout the post-monsoon recession."
)

h3("2.4.4 Channel routing methods and hydraulic coupling")
body(
    "Flow routing through channel reaches is available in HEC-HMS through the Muskingum, Muskingum\u2013Cunge, "
    "kinematic-wave, modified-Puls, and lag methods (USACE, n.d.). The Muskingum method, which conceptualises a reach as a "
    "linear storage element governed by a travel-time parameter and a dimensionless weighting parameter, remains the most "
    "commonly adopted routing approach in applications comparable to the present study (Hamdan et al., 2021; Labade et "
    "al., 2025; Nujhat et al., 2024). Its principal limitation, shared by all hydrologic routing schemes, is the inability "
    "to represent backwater effects, flow reversal, or looped stage\u2013discharge relationships, all of which are "
    "characteristic of extremely flat deltaic channels in which downstream water level rather than upstream inflow "
    "governs conveyance."
)
body(
    "Where such hydraulic controls are material, contemporary practice couples HEC-HMS with the Hydrologic Engineering "
    "Center River Analysis System (HEC-RAS). Zhang et al. (2022) developed an integrated flood-risk framework for the Old "
    "Brahmaputra River floodplain of Bangladesh in which HEC-HMS supplied inflow hydrographs to coupled one- and "
    "two-dimensional HEC-RAS models, achieving Nash\u2013Sutcliffe efficiency values of 0.93 and 0.81 and percent bias of "
    "\u22121.17% and 2.40% during calibration and validation, respectively. Admas et al. (2025) applied an equivalent "
    "coupling to the Gumara River in the Upper Blue Nile Basin to quantify inundation extents for return periods between 2 "
    "and 100 years and to evaluate the hydraulic effect of dyke construction, while Jawad (2024) employed the same "
    "architecture to generate flood inundation maps along the lowermost reach of the Brahmaputra. Because drainage of the "
    "Gumai Beel is regulated in part by water levels within the receiving Ichamoti channel rather than by catchment "
    "inflow alone, an analogous coupling represents the logical methodological extension of the present work."
)

h2("2.5 Integration of GIS and remote sensing in hydrological modelling")

h3("2.5.1 Digital elevation models and watershed delineation")
body(
    "Semi-distributed modelling is initiated through terrain analysis, in which sub-basin boundaries and stream networks "
    "are delineated from a digital elevation model (DEM) and physiographic attributes such as drainage area, mean slope, "
    "and longest flow path are extracted for each computational unit. A widening family of freely available global "
    "elevation products, including SRTM, NASADEM, ASTER, AW3D30, MERIT, and TanDEM-X, now competes for this role, and the "
    "differences among them are hydrologically consequential. Uuemaa et al. (2020), conducting a systematic multi-region "
    "accuracy assessment against LiDAR and Pleiades reference surfaces comprising 608 million pixels, established that "
    "vertical accuracy varies substantially among products and that terrain slope exerts the strongest control on error "
    "magnitude. AW3D30 exhibited the most stable performance across geographic settings, while NASADEM offered only "
    "marginal improvement relative to its SRTM predecessor."
)
body(
    "The reliability of automated delineation deteriorates precisely within the terrain class represented by the present "
    "study area. Datta et al. (2022), working on the Halda watershed of Bangladesh, demonstrated systematically that "
    "delineation outcomes are materially conditioned by the selection of DEM product, its spatial resolution, and the "
    "contributing-area threshold adopted for stream definition, and that these choices propagate into sub-basin geometry "
    "and all subsequently derived parameters. In low-relief deltaic terrain, where total elevation range may amount to "
    "only a few metres and where anthropogenic features such as roads, embankments, and sluice structures rather than "
    "natural topography govern actual flow pathways, DEM vertical error can exceed the topographic signal that the "
    "delineation algorithm is intended to detect. The literature therefore recommends systematic verification of "
    "automatically delineated drainage against hydrographic maps and field observation, a procedure adopted in the "
    "methodology of the present study."
)

h3("2.5.2 Land use, soil data, and curve number derivation")
body(
    "Parameterisation of the loss model in sparsely gauged basins depends upon thematic spatial mapping. The established "
    "workflow intersects a land-use and land-cover classification, commonly derived from Landsat or Sentinel-2 imagery, "
    "with a hydrologic soil group layer obtained from national soil surveys or global soil databases, and assigns "
    "area-weighted composite curve numbers to each sub-basin through standard lookup tabulations (Ranjan & Singh, 2022; "
    "Turkar et al., 2025). Ranjan and Singh (2022) implemented this procedure within ArcGIS for the Punpun River Basin, "
    "generating land-use, soil, and slope layers from which curve numbers were computed as direct HEC-HMS input, and "
    "reported coefficient of determination and Nash\u2013Sutcliffe efficiency values exceeding 0.75 for monthly and "
    "monsoonal models with percent bias below 10%."
)
body(
    "The sensitivity of simulated runoff to this parameterisation imposes a corresponding requirement for classification "
    "accuracy. Turkar et al. (2025) identified curve number among the four most sensitive parameters across the "
    "applications they reviewed, and M. B. Haque et al. (2024) demonstrated that the rainfall\u2013runoff relationship of "
    "the Halda catchment responded strongly to sub-basin curve numbers determined by land-cover classification. Remote "
    "sensing evidence further documents the rapidity of land-cover transformation within Bangladeshi wetland environments: "
    "Bhattacharjee et al. (2021), analysing multi-decadal Landsat imagery over a north-eastern wetland system, quantified "
    "substantial conversion of wetland and vegetation classes to agricultural and settlement uses. This consideration "
    "assumes particular importance for beel catchments, in which the land surface alternates seasonally among open water, "
    "cultivated land, and marsh vegetation, such that a single-date classification cannot adequately represent the "
    "hydrological character of the catchment across the annual cycle."
)

h3("2.5.3 Precipitation inputs in data-scarce catchments")
body(
    "Precipitation constitutes the dominant forcing variable of any rainfall\u2013runoff model, and its estimation "
    "represents the principal source of predictive uncertainty where rain-gauge networks are sparse or discontinuous. Two "
    "families of alternatives to gauge interpolation have reached operational maturity: satellite multi-sensor products, "
    "foremost the Integrated Multi-satellitE Retrievals for Global Precipitation Measurement (IMERG), and atmospheric "
    "reanalyses, foremost the fifth-generation European Centre for Medium-Range Weather Forecasts reanalysis (ERA5) "
    "described by Hersbach et al. (2020). Pradhan et al. (2022), synthesising the global body of IMERG validation "
    "studies, concluded that the product represents regional precipitation patterns and spatial means reliably and "
    "improves systematically with successive algorithm versions, while performing more accurately at monthly and annual "
    "aggregations than at daily and sub-daily resolutions and retaining recognised limitations for extreme intensities and "
    "orographically complex terrain."
)
body(
    "Operational experience confirms that satellite and reanalysis forcing can support credible hydrological simulation "
    "even in large, sparsely instrumented basins. Jawad (2024) forced coupled HEC-HMS and HEC-RAS models of the "
    "transboundary Brahmaputra Basin with near-real-time IMERG and GSMaP products, applying gauge-based correction where "
    "ground data permitted, and obtained acceptable discharge simulation and flood inundation mapping at Bahadurabad. "
    "Because the present study can draw upon Bangladesh Water Development Board and Bangladesh Meteorological Department "
    "observations within and adjacent to the Pabna region, this literature supports a strategy of gauge-primary forcing "
    "supplemented by reanalysis products for gap-filling and internal consistency verification."
)

h2("2.6 Model calibration, validation, and performance evaluation")
body(
    "Credible application of a rainfall\u2013runoff model requires a disciplined testing protocol in which model "
    "parameters are estimated against one portion of the observational record and predictive skill is demonstrated "
    "against an independent portion. This split-sample framework is followed, at least in its elementary form, by "
    "essentially all of the applications reviewed in this chapter (Hamdan et al., 2021; Nujhat et al., 2024; Ranjan & "
    "Singh, 2022; Zhang et al., 2022). The framework itself, however, has recently been subjected to systematic "
    "re-examination. Shen et al. (2022), conducting a large-sample experiment across 463 catchments, two conceptual "
    "hydrological models, and 50 alternative data-splitting schemes, demonstrated that the widespread practice of "
    "calibrating against older records and validating against more recent ones consistently degraded subsequent "
    "predictive performance, and that calibration against the full available record represented the most robust strategy "
    "for operational application. In a thesis context, where independent demonstration of predictive skill remains an "
    "examination requirement, these findings support the reporting of both a conventional split-sample evaluation and a "
    "final parameter set re-estimated across the complete record."
)
body(
    "Quantitative evaluation of model performance rests upon a compact set of statistical indicators. Review evidence "
    "indicates that Nash\u2013Sutcliffe efficiency, the coefficient of determination, and root-mean-square error "
    "constitute the most frequently reported metrics across HEC-HMS applications (Labade et al., 2025), with percent bias "
    "and the ratio of root-mean-square error to the standard deviation of observations increasingly reported alongside "
    "them (Ranjan & Singh, 2022). The interpretation of these indicators, however, is not mechanical. Althoff and "
    "Rodrigues (2021) established through systematic analysis that the goodness-of-fit criterion adopted as the objective "
    "function materially determines the resulting parameter set, and recommended that criteria be matched deliberately to "
    "the modelling purpose, with peak-weighted criteria appropriate to flood design and volume-oriented criteria "
    "appropriate to water-balance assessment. Their analysis implies that the reporting of a single aggregate efficiency "
    "score, detached from the objective function used to obtain it, provides an incomplete account of model performance."
)
body(
    "Calibration procedures range from manual adjustment guided by physical reasoning to fully automated optimisation. "
    "HEC-HMS provides univariate-gradient and simplex search algorithms in combination with a selection of objective "
    "functions (USACE, n.d.), and the reviewed applications span this full methodological range: manual calibration "
    "informed by parameter plausibility (Nujhat et al., 2024), automated optimisation using gradient search (Jawad, 2024), "
    "and staged procedures in which formal sensitivity analysis precedes calibration in order to concentrate effort upon "
    "the most influential parameters (Admas et al., 2025). The consistent identification of curve number, lag time, and "
    "routing parameters as the dominant controls upon model response (Turkar et al., 2025) provides a defensible "
    "starting parameter set for the present study. Reporting of calibrated parameter ranges alongside point estimates is "
    "further warranted, since alternative objective functions may select materially different yet equally acceptable "
    "parameter combinations (Althoff & Rodrigues, 2021)."
)

h2("2.7 Applications of HEC-HMS in international practice")
body(
    "The international application literature published since 2020 is extensive, and a selective review of "
    "methodologically instructive studies is presented here; Table 2.1 summarises their principal characteristics. Hamdan "
    "et al. (2021) developed a daily-timestep model of the semi-arid Al-Adhaim River catchment in northern Iraq using "
    "HEC-GeoHMS terrain preprocessing together with the SCS-CN, SCS unit hydrograph, and Muskingum configuration, "
    "calibrating across two hydrological years and verifying across a third. The study reported coefficients of "
    "determination of approximately 0.90 for both phases and concluded that the model was suitable for reservoir inflow "
    "estimation, although simulated dam discharge was slightly overestimated. Ranjan and Singh (2022) applied a "
    "comparable configuration to the Punpun River Basin of eastern India across the period 2005 to 2017, developing "
    "parallel daily, monthly, and monsoonal formulations and demonstrating that temporal aggregation materially improved "
    "performance, with the monthly model outperforming both alternatives."
)
body(
    "Aliye et al. (2020) contributed a comparative evaluation of HEC-HMS against SWAT for a data-scarce catchment of the "
    "Ethiopian Rift Valley Lakes Basin, establishing the viability of the lighter-parameterised framework in conditions of "
    "limited observational support. Admas et al. (2025) extended the Ethiopian experience toward integrated flood-risk "
    "analysis, coupling HEC-HMS design hydrographs with two-dimensional HEC-RAS hydraulics for the Gumara River floodplain "
    "and quantifying the reduction in inundated area achieved through dyke construction across return periods from 2 to "
    "100 years. Collectively, this literature supports three generalisations. First, the combination of SCS-CN loss, SCS "
    "unit hydrograph transform, and Muskingum routing constitutes the effective default configuration for event-scale "
    "application and performs satisfactorily across monsoonal and semi-arid regimes (Hamdan et al., 2021; Labade et al., "
    "2025). Second, the curve number is almost invariably the most sensitive parameter, such that its spatial estimation "
    "and calibration bounds govern overall model quality (Turkar et al., 2025). Third, model performance is constrained "
    "less by algorithmic structure than by the quality of input data, and by precipitation representation above all "
    "(Jawad, 2024; Pradhan et al., 2022)."
)

# ----------------------------------------------------------------- Table 2.1
tcap = doc.add_paragraph()
r = tcap.add_run("Table 2.1")
r.bold = True
tcap2 = doc.add_paragraph()
r = tcap2.add_run("Summary of Selected HEC-HMS Applications Reviewed in This Chapter")
r.italic = True

rows = [
    ("Study", "Study area", "Methods adopted", "Reported performance"),
    ("Aliye et al. (2020)", "Rift Valley Lakes Basin, Ethiopia",
     "HEC-HMS and SWAT comparison", "Both models applicable in data-scarce conditions"),
    ("Raihan et al. (2020)", "Upper Halda Basin, Bangladesh",
     "SWAT (comparative baseline)", "R\u00b2 = 0.80; NSE = 0.71"),
    ("Hamdan et al. (2021)", "Al-Adhaim catchment, Iraq",
     "SCS-CN; SCS UH; Muskingum; HEC-GeoHMS", "R\u00b2 \u2248 0.90 (calibration and verification)"),
    ("S. Haque et al. (2021)", "Brahmaputra Basin, Bangladesh",
     "Continuous simulation; MUSLE sediment routing", "NSE = 0.65 (cal.); 0.54 (val.)"),
    ("Ranjan & Singh (2022)", "Punpun River Basin, India",
     "SCS-CN; ArcGIS-derived CN; daily/monthly/monsoonal", "R\u00b2 and NSE > 0.75; PBIAS < 10%"),
    ("Zhang et al. (2022)", "Old Brahmaputra floodplain, Bangladesh",
     "HEC-HMS coupled with HEC-RAS 1D/2D", "NSE = 0.93 (cal.); 0.81 (val.)"),
    ("Nur et al. (2022)", "Khowai River Basin, Bangladesh",
     "Event-based HEC-HMS simulation", "Model applicable to flash-flood-prone basin"),
    ("Jawad (2024)", "Brahmaputra Basin (transboundary)",
     "SMA; Clark UH; linear reservoir; IMERG/GSMaP; HEC-RAS", "Credible discharge and inundation from satellite forcing"),
    ("Nujhat et al. (2024)", "Gumti River Basin, Bangladesh",
     "SCS-CN; Muskingum; SRTM delineation", "R\u00b2 = 0.64 (cal.); 0.68 (val.); PBIAS very good"),
    ("M. B. Haque et al. (2024)", "Halda River catchment, Bangladesh",
     "SCS-CN optimised against SWAT-derived values", "NSE = 0.72 (cal.); 0.82 (val.); baseflow underestimated"),
    ("Admas et al. (2025)", "Gumara River, Upper Blue Nile Basin",
     "HEC-HMS coupled with HEC-RAS; return-period analysis", "Inundation quantified for 2\u2013100-year events"),
    ("Moniruzzaman & Mahalder (2026)", "Atrai\u2013Karatoa River Basin, Bangladesh",
     "HEC-HMS and SWAT comparison", "HEC-HMS R\u00b2 = 0.70 (cal.); 0.56 (val.)"),
]

table = doc.add_table(rows=len(rows), cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Inches(1.5), Inches(1.7), Inches(1.75), Inches(1.55)]
for i, row_data in enumerate(rows):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.width = widths[j]
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(cell_text)
        run.font.size = Pt(10)
        run.font.name = FONT
        if i == 0:
            run.bold = True

note = doc.add_paragraph()
r = note.add_run(
    "Note. NSE = Nash\u2013Sutcliffe efficiency; R\u00b2 = coefficient of determination; PBIAS = percent bias; "
    "CN = curve number; UH = unit hydrograph; SMA = soil moisture accounting; cal. = calibration; val. = validation."
)
r.italic = True
r.font.size = Pt(10)

h2("2.8 Applications of HEC-HMS in Bangladesh")
body(
    "Bangladesh occupies the terminal reach of the Ganges\u2013Brahmaputra\u2013Meghna system, and its hydrological "
    "regime is characterised by intense monsoonal rainfall, transboundary inflow, extremely low topographic relief, and an "
    "extensive network of floodplain depressions. Within this setting, the national academic literature has increasingly "
    "adopted HEC-HMS for basin-scale rainfall\u2013runoff investigation, and the record published since 2020 now spans the "
    "principal basin types of the country."
)
body(
    "S. Haque et al. (2021) developed a continuous HEC-HMS model of the poorly gauged Brahmaputra Basin, calibrated "
    "against daily runoff at Bahadurabad for the period 1983 to 1996 and validated for 1997 to 2010, and extended the "
    "framework with the Modified Universal Soil Loss Equation and Engelund\u2013Hansen sediment routing in order to "
    "project future sediment load under the RCP8.5 scenario. Their results indicated increases in mean annual sediment "
    "load of 34%, 67%, and 115% by the 2020s, 2050s, and 2080s respectively, demonstrating the utility of HEC-HMS as the "
    "hydrological component of climate-impact assessment chains. Zhang et al. (2022) coupled HEC-HMS with one- and "
    "two-dimensional HEC-RAS models to construct an integrated flood risk, hazard, and vulnerability assessment for the "
    "Old Brahmaputra River floodplain, achieving Nash\u2013Sutcliffe efficiency values of 0.93 and 0.81 during calibration "
    "and validation and demonstrating that coupled hydrologic\u2013hydraulic frameworks can support the development of "
    "early-warning systems and adaptation strategies."
)
body(
    "At the sub-basin scale, Nur et al. (2022) applied HEC-HMS to rainfall\u2013runoff simulation of the Khowai River "
    "Basin, a flash-flood-prone catchment of north-eastern Bangladesh. Nujhat et al. (2024) calibrated and validated a "
    "model of the Gumti River Basin using observations from 2019 to 2021 with SRTM-based catchment delineation, Muskingum "
    "routing, and SCS-CN losses, obtaining coefficients of determination of 0.64 and 0.68 for calibration and validation "
    "respectively with percent bias in the very good performance category, and recommended the calibrated model as a "
    "planning instrument for flood prediction and water-resources management. M. B. Haque et al. (2024) modelled the "
    "ecologically significant Halda catchment, optimising curve numbers against values derived from a companion SWAT "
    "application and achieving Nash\u2013Sutcliffe efficiency values of 0.72 and 0.82, while explicitly documenting "
    "underestimation of baseflow during the dry season. Moniruzzaman and Mahalder (2026) compared HEC-HMS with SWAT for "
    "the Atrai\u2013Karatoa River Basin of northern Bangladesh, a system hydrologically contiguous with the greater Chalan "
    "Beel floodplain to which the Pabna beel network belongs, reporting that SWAT captured high flows more accurately "
    "while HEC-HMS performed better across medium flows, with HEC-HMS achieving coefficients of determination of 0.70 and "
    "0.56 during calibration and validation."
)
body(
    "Complementary investigations employing alternative modelling frameworks complete the national picture. Raihan et al. "
    "(2020) simulated streamflow in the Upper Halda Basin using SWAT, achieving a coefficient of determination of 0.80 and "
    "Nash\u2013Sutcliffe efficiency of 0.71, and identified groundwater delay time, baseflow alpha factor, and curve number "
    "as the parameters exerting greatest influence upon model performance while emphasising the difficulty of representing "
    "spatial rainfall variability from a single gauging station. Akter and Sawon (2024) developed hydrological models of "
    "selected flash-flood-prone rivers of Bangladesh, contributing to the understanding of rapid-response catchment "
    "behaviour under monsoonal forcing. Datta et al. (2022) provided the DEM sensitivity analysis of watershed delineation "
    "discussed in Section 2.5.1, and Jawad (2024) demonstrated satellite-forced coupled hydrologic\u2013hydraulic modelling "
    "for the Bangladeshi reach of the Brahmaputra."
)
body(
    "Considered collectively, the Bangladeshi HEC-HMS literature exhibits a pronounced spatial concentration. Published "
    "applications cluster within large transboundary basins, flashy piedmont and hill catchments, and the "
    "Atrai\u2013Karatoa corridor. The moribund distributary systems of the Ganges right bank, of which the Ichamoti River "
    "of Pabna constitutes a representative example, together with their associated beel catchments, have not been the "
    "subject of any published HEC-HMS rainfall\u2013runoff investigation. It is precisely within such systems, where local "
    "rainfall accumulation and drainage congestion rather than upstream flood waves govern inundation behaviour, that "
    "quantification of runoff is most necessary for drainage design, and yet where the observational network is thinnest."
)

h2("2.9 Wetland and beel hydrology in the Bangladesh floodplain")
body(
    "Wetland hydrology differs fundamentally from upland catchment hydrology in that the water regime, understood as the "
    "seasonal pattern of water-level rise, persistence, and recession, functions as the master variable governing "
    "biogeochemical processes, ecological structure, and land-use viability. Altered hydrology is correspondingly "
    "identified as the principal pathway through which environmental and anthropogenic change degrades wetland function, "
    "since modification of water depth and residence time alters the balance between organic-matter production and "
    "decomposition and may convert a wetland from a net sink to a net source of carbon and nutrients (Salimi et al., "
    "2021). Within the floodplain morphology of Bangladesh, the characteristic wetland landform is the beel, a "
    "saucer-shaped depression retaining water perennially or seasonally and receiving inflow both from direct "
    "rainfall\u2013runoff generated within its local catchment and from lateral spill originating in adjacent river "
    "channels during the monsoon."
)
body(
    "The hydrological functioning of these depressions has been substantially modified by embankment and polder "
    "infrastructure constructed from the 1960s onward. Adnan et al. (2020), analysing the embanked delta of south-western "
    "Bangladesh, described the systemic mechanism through which such enclosure generates internal drainage congestion: "
    "embankments sever the hydraulic connection between floodplain and river channel, sediment accumulates within "
    "riverbeds rather than upon the floodplain surface, riverbed elevation rises relative to the enclosed land, gravity "
    "drainage through sluice structures progressively fails, and monsoon runoff accumulates within the beels as pluvial "
    "flooding. Their analysis established that a large majority of agricultural and aquacultural land within the embanked "
    "region now lies within flood-susceptible zones, and evaluated controlled sediment reintroduction across 106 candidate "
    "beels as a rehabilitation strategy capable of raising land elevation by up to 1.4 m over five years. The mechanism "
    "described is structural and regional rather than locally anomalous, and it applies with equal force to the embanked "
    "interior floodplains of the Ganges right bank."
)
body(
    "A second control operates at basin scale through the long-term reduction of dry-season flow within the Ganges "
    "distributary network. Ali and Hasan (2022), analysing Bangladesh Water Development Board discharge records for the "
    "Gorai River, the principal Ganges distributary and the closest regional analogue to the Ichamoti, determined that "
    "mean annual flow during 2000 to 2016 was approximately 13% lower than during 1984 to 1999, that mean monthly high "
    "flow declined by 20%, and that the river now frequently fails to satisfy an environmental flow requirement estimated "
    "at 295 m\u00b3/s, or 29% of mean annual flow, with a deficient-flow condition persisting from December through May. "
    "Reduced discharge within parent channels deprives distributary offtakes of the sediment-flushing flows required to "
    "maintain conveyance capacity, initiating progressive siltation and eventual hydraulic disconnection of distributary "
    "systems."
)
body(
    "The consequences of these combined mechanisms are documented directly for the wetland complex surrounding the study "
    "area. Yankyera and Alam (2025), analysing five decades of Landsat imagery over Chalan Beel, the largest wetland "
    "ecosystem of Bangladesh and the system into which the Pabna floodplain drains, identified a consistent trajectory of "
    "expanding farmland and developed land accompanied by declining vegetation and wetland extent, with the most "
    "pronounced transformation occurring between 2013 and 2023. Bhattacharjee et al. (2021) reported an equivalent "
    "trajectory for a north-eastern wetland ecosystem. These transformations are hydrologically significant because "
    "reduction in wetland storage capacity diminishes the attenuation that beels provide to monsoon runoff, thereby "
    "increasing peak discharge and inundation depth within the residual drainage network."
)

h2("2.10 The Ichamoti River and Gumai Beel system")
body(
    "The Ichamoti River of Pabna District exemplifies the condition described in the preceding section. The river "
    "originates from the Padma in the vicinity of Shibrampur within Pabna Sadar Upazila and follows a course of "
    "approximately 82 to 84 km through Pabna town toward the Hurasagar system in Bera Upazila. Over recent decades the "
    "channel has lost effective hydraulic connectivity with its parent rivers and has become progressively constricted "
    "through siltation, encroachment, solid-waste disposal, and aquatic weed infestation. Parvez et al. (2021), "
    "investigating the causes of waterlogging within Pabna Municipality, documented that the Ichamoti bisects the "
    "municipality but no longer functions as an effective drainage outfall, being choked with water weeds and accumulated "
    "sediment, and reported that residents identified blocked and undersized drains, absent operation and maintenance "
    "provision, and an unplanned drainage network as the leading causes of inundation following heavy rainfall. Their "
    "physical inventory recorded approximately 7.9 km of the Ichamoti channel within the municipal area alongside 11.29 km "
    "of primary drains and 56.23 km of secondary drains, establishing the scale of the drainage system dependent upon the "
    "river as outfall."
)
body(
    "Recognition of this condition has prompted large-scale government intervention. A rejuvenation programme valued at "
    "Tk 1,554.90 crore, implemented under the Bangladesh Water Development Board, provides for dredging across a 33.77 km "
    "reach of the river together with 44.07 km of connecting canals, 20 km of the Sutikhali River, and 12.37 km of the "
    "Bharara channel, with the stated objective of restoring hydraulic connectivity between the Ichamoti and the Padma and "
    "Jamuna systems (\u201cTk 1,554cr Project,\u201d 2024). The hydrological design basis for such intervention requires "
    "quantitative estimation of the runoff volumes and hydrograph characteristics that the restored corridor must convey, "
    "which is precisely the information that a calibrated rainfall\u2013runoff model provides."
)
body(
    "The Gumai Beel, which constitutes the focus of the present investigation, forms one component of the beel network of "
    "the Pabna floodplain draining through the Ichamoti corridor. Its inundation behaviour is governed by the interaction "
    "of three controls identified in the reviewed literature as characteristic of embanked distributary floodplains: local "
    "monsoon rainfall\u2013runoff generated within the beel catchment; the conveyance capacity of the silted Ichamoti "
    "channel and its associated khals, which determines the rate at which stored water may be evacuated; and the stage of "
    "the receiving river system, which may impose backwater constraints upon gravity drainage (Adnan et al., 2020; Parvez "
    "et al., 2021). A rainfall\u2013runoff model of the beel catchment constitutes the necessary first element of any "
    "quantitative analysis of this system, since it supplies the inflow boundary condition required for drainage design, "
    "for evaluation of the hydrological benefits attributable to the rejuvenation programme, and for assessment of "
    "waterlogging risk under current and projected rainfall regimes."
)

h2("2.11 Research gap and summary")
body(
    "The literature reviewed in this chapter supports four conclusions that collectively establish the position of the "
    "present study. First, conceptual semi-distributed modelling using HEC-HMS constitutes a mature and extensively "
    "validated approach whose conventional method combination of SCS-CN loss, SCS unit hydrograph transform, and "
    "Muskingum routing performs satisfactorily across monsoonal and semi-arid regimes, provided that method selection is "
    "justified with reference to catchment characteristics and modelling purpose and that curve number estimation receives "
    "particular methodological attention (Hamdan et al., 2021; Horton et al., 2022; Labade et al., 2025; Soulis, 2021). "
    "Second, the spatial data infrastructure required for such modelling, comprising global elevation products, satellite "
    "land-cover classification, and satellite or reanalysis precipitation, is available for Bangladesh, but low-relief "
    "deltaic terrain imposes recognised constraints upon automated delineation and satellite precipitation products retain "
    "documented limitations for extreme intensities, both of which require deliberate management (Datta et al., 2022; "
    "Pradhan et al., 2022; Uuemaa et al., 2020)."
)
body(
    "Third, published Bangladeshi applications of HEC-HMS concentrate within large transboundary basins, flashy piedmont "
    "catchments, and the Atrai\u2013Karatoa corridor (Jawad, 2024; M. B. Haque et al., 2024; Moniruzzaman & Mahalder, "
    "2026; Nujhat et al., 2024; Nur et al., 2022; S. Haque et al., 2021; Zhang et al., 2022). No published study has "
    "developed a rainfall\u2013runoff model for the beel catchments of the moribund Ganges distributaries, of which the "
    "Ichamoti\u2013Gumai Beel system is representative. Fourth, the hydrological problem posed by such systems is "
    "distinctive in kind rather than merely in degree: inundation is governed by locally generated runoff accumulating "
    "behind congested drainage rather than by riverine flood waves propagating from upstream (Adnan et al., 2020; Parvez "
    "et al., 2021), while the long-term deterioration of the distributary network reflects basin-scale flow reduction that "
    "has been quantified for the neighbouring Gorai but has never been connected to a quantitative runoff model of the "
    "beel itself (Ali & Hasan, 2022)."
)
body(
    "The research gap addressed by the present study follows directly from these conclusions. A calibrated and validated "
    "HEC-HMS rainfall\u2013runoff model of the Gumai Beel catchment will provide the first quantitative estimate of the "
    "runoff volumes and hydrograph dynamics that the Ichamoti drainage corridor is required to convey, thereby "
    "establishing the hydrological foundation for drainage design and for evaluation of the ongoing river rejuvenation "
    "programme. Methodologically, the study extends the Bangladeshi HEC-HMS literature into an ultra-flat, "
    "wetland-dominated, and data-scarce environment in which DEM limitations, seasonal alternation of land cover, and "
    "wetland storage behaviour present challenges that have been identified but not resolved in previous work. The methods "
    "adopted in the following chapter, comprising comparative evaluation of loss and transform formulations, "
    "field-verified catchment delineation, gauge-primary precipitation forcing supplemented by reanalysis data, "
    "split-sample testing informed by current calibration research, and performance evaluation using purpose-matched "
    "objective functions together with multiple statistical indicators, are each grounded in the literature reviewed in "
    "this chapter (Althoff & Rodrigues, 2021; Shen et al., 2022; Turkar et al., 2025)."
)

# ================================================================ REFERENCES
page_break()
h1("References")

references = [
    "Admas, M., Asrade, T. M., & Cherie, W. D. (2025). Application of the HEC-RAS and HEC-HMS models for flood risk "
    "analysis in the Gumara River, Upper Blue Nile Basin, Ethiopia. Advances in Meteorology, 2025, Article 5092932. "
    "https://doi.org/10.1155/adme/5092932",

    "Adnan, M. S. G., Talchabhadel, R., Nakagawa, H., & Hall, J. W. (2020). The potential of Tidal River Management for "
    "flood alleviation in South Western Bangladesh. Science of the Total Environment, 731, Article 138747. "
    "https://doi.org/10.1016/j.scitotenv.2020.138747",

    "Akter, A., & Sawon, F. S. (2024). Hydrological modeling of the selected flash flood-prone rivers. Natural Hazards. "
    "https://doi.org/10.1007/s11069-024-06928-z",

    "Ali, M. S., & Hasan, M. M. (2022). Environmental flow assessment of Gorai River in Bangladesh: A comparative "
    "analysis of different hydrological methods. Heliyon, 8(7), Article e09857. "
    "https://doi.org/10.1016/j.heliyon.2022.e09857",

    "Aliye, M. A., Aga, A. O., Tadesse, T., & Yohannes, P. (2020). Evaluating the performance of HEC-HMS and SWAT "
    "hydrological models in simulating the rainfall-runoff process for data scarce region of Ethiopian Rift Valley Lake "
    "Basin. Open Journal of Modern Hydrology, 10(4), 105\u2013122. https://doi.org/10.4236/ojmh.2020.104007",

    "Althoff, D., & Rodrigues, L. N. (2021). Goodness-of-fit criteria for hydrological models: Model calibration and "
    "performance assessment. Journal of Hydrology, 600, Article 126674. https://doi.org/10.1016/j.jhydrol.2021.126674",

    "Bhattacharjee, S., Islam, M. T., Kabir, M. E., & Kabir, M. M. (2021). Land-use and land-cover change detection in a "
    "north-eastern wetland ecosystem of Bangladesh using remote sensing and GIS techniques. Earth Systems and "
    "Environment, 5(2), 319\u2013340. https://doi.org/10.1007/s41748-021-00228-3",

    "Datta, S., Karmakar, S., Mezbahuddin, S., Chaudhary, B. S., Hossain, M. M., Hoque, M. E., Abdullah-Al-Mamun, M. M., "
    "& Baul, T. K. (2022). The limits of watershed delineation: Implications of different DEMs, DEM resolutions, and "
    "area threshold values. Hydrology Research, 53(8), 1047\u20131062. https://doi.org/10.2166/nh.2022.126",

    "Dibaba, W. T., Demissie, T. A., & Miegel, K. (2020). Watershed hydrological response to combined land use/land cover "
    "and climate change in highland Ethiopia: Finchaa catchment. Water, 12(6), Article 1801. "
    "https://doi.org/10.3390/w12061801",

    "Hamdan, A. N. A., Almuktar, S., & Scholz, M. (2021). Rainfall-runoff modeling using the HEC-HMS model for the "
    "Al-Adhaim River catchment, northern Iraq. Hydrology, 8(2), Article 58. https://doi.org/10.3390/hydrology8020058",

    "Haque, M. B., Karmakar, S., & Hossain, M. M. (2024). Rainfall-runoff modeling using the HEC-HMS flow modeling "
    "framework for the Halda River catchment, Bangladesh [Preprint]. Research Square. "
    "https://doi.org/10.21203/rs.3.rs-3824469/v1",

    "Haque, S., Ali, M. M., Islam, A. K. M. S., & Khan, M. J. U. (2021). Changes in flow and sediment load of poorly "
    "gauged Brahmaputra river basin under an extreme climate scenario. Journal of Water and Climate Change, 12(3), "
    "937\u2013954. https://doi.org/10.2166/wcc.2020.219",

    "Hersbach, H., Bell, B., Berrisford, P., Hirahara, S., Hor\u00e1nyi, A., Mu\u00f1oz-Sabater, J., Nicolas, J., Peubey, "
    "C., Radu, R., Schepers, D., Simmons, A., Soci, C., Abdalla, S., Abellan, X., Balsamo, G., Bechtold, P., Biavati, G., "
    "Bidlot, J., Bonavita, M., \u2026 Th\u00e9paut, J.-N. (2020). The ERA5 global reanalysis. Quarterly Journal of the "
    "Royal Meteorological Society, 146(730), 1999\u20132049. https://doi.org/10.1002/qj.3803",

    "Horton, P., Schaefli, B., & Kauzlaric, M. (2022). Why do we have so many different hydrological models? A review "
    "based on the case of Switzerland. WIREs Water, 9(1), Article e1574. https://doi.org/10.1002/wat2.1574",

    "Jawad, M. (2024). Evaluation of near real-time Global Precipitation Measurement (GPM) precipitation products for "
    "hydrological modelling and flood inundation mapping of sparsely gauged large transboundary basins\u2014A case study "
    "of the Brahmaputra basin. Remote Sensing, 16(10), Article 1756. https://doi.org/10.3390/rs16101756",

    "Labade, P., Ayare, B. L., Bhange, H. N., Ingle, P. M., & Kolhe, P. R. (2025). A comprehensive review on "
    "rainfall-runoff modelling using HEC-HMS. International Journal of Research in Agronomy, 8(10), 1130\u20131138. "
    "https://doi.org/10.33545/2618060X.2025.v8.i10o.4374",

    "Moniruzzaman, M., & Mahalder, B. (2026). Assessing SWAT and HEC-HMS model efficiency for watershed management in "
    "the Atrai-Karatoa River Basin, Bangladesh. Evolving Earth, 4, Article 100136. "
    "https://doi.org/10.1016/j.eve.2026.100136",

    "Nujhat, M., Rayhan, M., & Amin, M. K. (2024). Hydrological modelling and its implication in sustainable water "
    "resource management in Gumti River Basin in Bangladesh. International Journal of Sustainability in Energy and "
    "Environment, 1(2), 40\u201348.",

    "Nur, F., Mohib, K. M., Toma, U. T., Rozario, P. M., Bari, M. S., Anjum, N., & Khandakar, F. (2022, February 10\u201312). "
    "Rainfall-runoff simulation of Khowai River basin using HEC-HMS model [Conference presentation]. 6th International "
    "Conference on Civil Engineering for Sustainable Development (ICCESD 2022), Khulna University of Engineering & "
    "Technology, Khulna, Bangladesh.",

    "Parvez, M., Sadat, N., Tasnim, F., & Nejhum, I. J. (2021). Identifying the causes of waterlogging on people\u2019s "
    "perception towards a resilient community: A case study on Pabna Municipality, Bangladesh. Ecofeminism and Climate "
    "Change, 2(3), 110\u2013126. https://doi.org/10.1108/EFCC-11-2020-0033",

    "Pradhan, R. K., Markonis, Y., Vargas Godoy, M. R., Villalba-Pradas, A., Andreadis, K. M., Nikolopoulos, E. I., "
    "Papalexiou, S. M., Rahim, A., Tapiador, F. J., & Hanel, M. (2022). Review of GPM IMERG performance: A global "
    "perspective. Remote Sensing of Environment, 268, Article 112754. https://doi.org/10.1016/j.rse.2021.112754",

    "Raihan, F., Beaumont, L. J., Maina, J., Saiful Islam, A. K. M., & Harrison, S. P. (2020). Simulating streamflow in "
    "the Upper Halda Basin of southeastern Bangladesh using SWAT model. Hydrological Sciences Journal, 65(1), "
    "138\u2013151. https://doi.org/10.1080/02626667.2019.1682149",

    "Ranjan, S., & Singh, V. P. (2022). HEC-HMS based rainfall-runoff model for Punpun river basin. Water Practice & "
    "Technology, 17(5), 986\u20131001. https://doi.org/10.2166/wpt.2022.033",

    "Sahu, M. K., Shwetha, H. R., & Dwarakish, G. S. (2023). State-of-the-art hydrological models and application of the "
    "HEC-HMS model: A review. Modeling Earth Systems and Environment, 9(3), 3029\u20133051. "
    "https://doi.org/10.1007/s40808-023-01704-7",

    "Salimi, S., Almuktar, S. A. A. A. N., & Scholz, M. (2021). Impact of climate change on wetland ecosystems: A "
    "critical review of experimental wetlands. Journal of Environmental Management, 286, Article 112160. "
    "https://doi.org/10.1016/j.jenvman.2021.112160",

    "Shen, H., Tolson, B. A., & Mai, J. (2022). Time to update the split-sample approach in hydrological model "
    "calibration. Water Resources Research, 58(3), Article e2021WR031523. https://doi.org/10.1029/2021WR031523",

    "Shi, W., & Wang, N. (2020). An improved SCS-CN method incorporating slope, soil moisture, and storm duration factors "
    "for runoff prediction. Water, 12(5), Article 1335. https://doi.org/10.3390/w12051335",

    "Soulis, K. X. (2021). Soil Conservation Service Curve Number (SCS-CN) method: Current applications, remaining "
    "challenges, and future perspectives. Water, 13(2), Article 192. https://doi.org/10.3390/w13020192",

    "Tk 1,554cr project to revive dying Ichamati. (2024, November 26). The Daily Star. "
    "https://www.thedailystar.net/news/bangladesh/news/tk-1554cr-project-revive-dying-ichamati-3555931",

    "Turkar, R. K., Shrivastava, R. N., Awasthi, M. K., & Lodhi, A. S. (2025). A comprehensive review of the HEC-HMS "
    "rainfall\u2013runoff simulation model and its hydrological applications. Journal of Agriculture and Ecology Research "
    "International, 26(5), 130\u2013144. https://doi.org/10.9734/jaeri/2025/v26i5711",

    "U.S. Army Corps of Engineers, Hydrologic Engineering Center. (n.d.). HEC-HMS technical reference manual. Retrieved "
    "August 17, 2026, from https://www.hec.usace.army.mil/confluence/hmsdocs/hmstrm",

    "Uuemaa, E., Ahi, S., Montibeller, B., Muru, M., & Kmoch, A. (2020). Vertical accuracy of freely available global "
    "digital elevation models (ASTER, AW3D30, MERIT, TanDEM-X, SRTM, and NASADEM). Remote Sensing, 12(21), Article 3482. "
    "https://doi.org/10.3390/rs12213482",

    "Yankyera, S., & Alam, B. M. (2025). Five decades of transformation due to human-environment stressors: Land cover, "
    "vegetation, and land surface temperature change analysis in the largest wetland ecosystem in Bangladesh. Earth "
    "Systems and Environment, 9(2), 589\u2013604. https://doi.org/10.1007/s41748-025-00652-9",

    "Zhang, K., Shalehy, M. H., Ezaz, G. T., Chakraborty, A. K., Mohib, K. M., & Liu, L. (2022). An integrated flood risk "
    "assessment approach based on coupled hydrological-hydraulic modeling and bottom-up hazard vulnerability analysis. "
    "Environmental Modelling & Software, 148, Article 105279. https://doi.org/10.1016/j.envsoft.2021.105279",
]

for r_text in references:
    ref(r_text)

OUT = "thesis/Literature_Review_Gumai_Beel_Ichamoti_HEC-HMS.docx"
doc.save(OUT)

words = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"Saved {OUT}")
print(f"Approximate word count (paragraphs only): {words}")
